"""
DOCX Extractor - Extract text from Word documents.

Strategy:
- Direct text extraction from paragraphs
- Tables extraction with formatting
- Simple and fast
"""

import logging
from typing import BinaryIO

from docx import Document

from domain.models.document import (
    ExtractionResult,
    Page,
    ExtractionMetadata,
    DocumentType
)
from domain.models.config import ExtractionConfig
from domain.exceptions import ExtractionError

logger = logging.getLogger(__name__)


class DOCXExtractor:
    """
    DOCX document extractor.

    Uses python-docx for direct text extraction.
    """

    @property
    def name(self) -> str:
        return "DOCX Extractor"

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return ('.docx', '.DOCX')

    def can_handle(self, file_name: str) -> bool:
        """Check if file is a DOCX"""
        return file_name.lower().endswith('.docx')

    def extract(
        self,
        file: BinaryIO,
        file_name: str,
        config: ExtractionConfig
    ) -> ExtractionResult:
        """
        Extract text from DOCX.

        Args:
            file: DOCX file object
            file_name: Original file name
            config: Extraction configuration

        Returns:
            ExtractionResult with extracted content

        Raises:
            ExtractionError: If extraction fails
        """
        import time
        start_time = time.time()

        try:
            file.seek(0)
            logger.info(f"Extracting DOCX: {file_name}")

            doc = Document(file)
            sections = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    sections.append(para.text)

            # Extract tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(cells))

                if table_rows:
                    sections.append("\n".join(table_rows))

            full_text = "\n".join(sections)

            processing_time = time.time() - start_time

            metadata = ExtractionMetadata(
                document_type=DocumentType.DOCX,
                pages_count=1,
                extraction_method="python-docx",
                processing_time_seconds=processing_time,
                file_size_bytes=self._get_file_size(file)
            )

            logger.info(
                f"DOCX extraction complete: {file_name} | "
                f"{len(sections)} sections | {processing_time:.2f}s"
            )

            return ExtractionResult(
                file_name=file_name,
                pages=[Page(number=1, text=full_text)],
                metadata=metadata
            )

        except Exception as e:
            logger.exception(f"DOCX extraction failed: {file_name}")
            raise ExtractionError(
                f"Failed to extract DOCX: {e}",
                file_name=file_name
            ) from e

    @staticmethod
    def _get_file_size(file: BinaryIO) -> int:
        """Get file size in bytes"""
        current_pos = file.tell()
        file.seek(0, 2)
        size = file.tell()
        file.seek(current_pos)
        return size
